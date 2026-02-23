import os
import glob
from docx import Document
from dotenv import load_dotenv
from supabase_rag import SupabaseMemory

def main():
    load_dotenv(override=True)
    
    api_key = os.environ.get("OPENAI_API_KEY")
    supabase_url = os.environ.get("SUPABASE_URL")
    supabase_key = os.environ.get("SUPABASE_KEY")
    
    import sys
    sys.stdout = open('debug_utf8.txt', 'w', encoding='utf-8')
    sys.stderr = sys.stdout

    memory = SupabaseMemory(
        openai_api_key=api_key,
        supabase_url=supabase_url,
        supabase_key=supabase_key
    )
    
    pops_dir = os.path.join(os.path.dirname(__file__), "data", "pops")
    docx_files = glob.glob(os.path.join(pops_dir, "*.docx"))
    
    for file_path in docx_files:
        filename = os.path.basename(file_path)
        print(f"Loading {filename}...")
        
        try:
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
                            
            text_content = "\n".join(full_text)
            print(f"File {filename} loaded. Size: {len(text_content)} chars")
            
            memory.ingest(text_content, source=filename)
        except Exception as e:
            print(f"Error processing {filename}: {e}")
            import traceback
            traceback.print_exc()

if __name__ == "__main__":
    main()
