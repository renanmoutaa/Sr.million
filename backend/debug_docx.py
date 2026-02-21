from docx import Document
from pathlib import Path
import os

# Absolute path based on verify location
POP_FILE = Path(r"d:\Sr. Million\backend\data\pops\POP 01 - Comercial - Aplicação Técnica.docx")

print(f"Checking {POP_FILE}...")
if not POP_FILE.exists():
    print("File does not exist!")
    exit(1)

try:
    doc = Document(POP_FILE)
    print(f"Document opened. Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")
    
    text = []
    # Paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            text.append(p.text)
            
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in table.rows[0].cells: # iterate first row to see structure? No, iterate all
                pass
            # Just extract text from cells
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text.append(" | ".join(row_text))

    full_text = "\n".join(text)
    print(f"Text length: {len(full_text)}")
    print("First 200 chars:")
    print(full_text[:200])
except Exception as e:
    print(f"Error opening/reading docx: {e}")
