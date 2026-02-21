import os
from docx import Document
import json
from pathlib import Path
from memory_rag import LightMemory
# Paths to the specific files
# ...

SETTINGS_FILE = Path('data/settings.json')

def load_settings():
    if not SETTINGS_FILE.exists():
        return {}
    with open(SETTINGS_FILE, 'r') as f:
        return json.load(f)
# Directory for POP files
BASE_DIR = Path(__file__).parent
POPS_DIR = BASE_DIR / 'data' / 'pops'

def get_pop_files():
    if not POPS_DIR.exists():
        print(f"Directory not found: {POPS_DIR}")
        return []
    return list(POPS_DIR.glob("*.docx"))

def read_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []
        
        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = []
                    for para in cell.paragraphs:
                        if para.text.strip():
                            cell_text.append(para.text)
                    if cell_text:
                        row_data.append(" ".join(cell_text))
                
                # Join row cells with separator
                if row_data:
                    # Remove duplicates if cell merging causes repetition?
                    # python-docx handles merged cells by repeating them.
                    # Simple strategy: keep unique
                    unique_data = []
                    seen = set()
                    for d in row_data:
                        if d not in seen:
                            unique_data.append(d)
                            seen.add(d)
                    full_text.append(" | ".join(unique_data))
                    
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return None

def main():
    print("Starting ingestion process...")
    
    # Load settings to get API Key
    settings = load_settings()
    api_key = settings.get("openai_api_key")
    
    if not api_key:
        print("Error: OpenAI API Key not found in settings.json.")
        return

    # Initialize Memory
    try:
        memory = LightMemory(openai_api_key=api_key)
        print("LightMemory initialized.")
    except Exception as e:
        print(f"Failed to init memory: {e}")
        return

    # Ingest Files
    # Ingest Files
    pop_files = get_pop_files()
    if not pop_files:
        print("No POP files found to ingest.")
        return

    for file_path in pop_files:
        path = file_path
        
        print(f"Reading {path.name}...")
        text = read_docx(path)
        
        if text:
            print(f"Ingesting {path.name} ({len(text)} chars)...")
            try:
                memory.ingest(text, source=path.name)
                print(f"Successfully ingested {path.name}")
            except Exception as e:
                print(f"Error ingesting {path.name}: {e}")
        else:
            print(f"Skipping {path.name} (empty or error)")

    print("All ingestions completed.")

if __name__ == "__main__":
    main()
